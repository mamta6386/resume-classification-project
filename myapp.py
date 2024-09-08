import streamlit as st
import re
import nltk
import pandas as pd
from nltk.tokenize import RegexpTokenizer
from nltk.stem import WordNetLemmatizer
from nltk.corpus import stopwords
from pdfminer.pdfdocument import PDFSyntaxError
from docx import Document
import io
import pickle
import pdfplumber

# Download NLTK resources
nltk.download('wordnet')
nltk.download('stopwords')

# Load the saved Naive Bayes model
with open('modelNB.pkl', 'rb') as file:
    loaded_nb_model = pickle.load(file)

def preprocess(sentence):
    sentence = str(sentence)
    sentence = sentence.lower()
    sentence = sentence.replace('{html}', "")
    cleanr = re.compile('<.*?>')
    cleantext = re.sub(cleanr, '', sentence)
    rem_url = re.sub(r'http\S+', '', cleantext)
    rem_num = re.sub('[0-9]+', '', rem_url)
    rem_email = re.sub(r'\S+@\S+', '', rem_num)
    tokenizer = RegexpTokenizer(r'\w+')
    tokens = tokenizer.tokenize(rem_email)
    stop = set(stopwords.words('english'))
    filtered_words = [w for w in tokens if len(w) > 2 if not w in stop]
    lemmatizer = WordNetLemmatizer()
    lemma_words = [lemmatizer.lemmatize(w) for w in filtered_words]
    return " ".join(lemma_words)

def extract_text_from_doc(doc_file):
    try:
        if doc_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":  # .docx
            resume_text = extract_text_from_docx(doc_file)
        elif doc_file.type == "application/msword":  # .doc (Microsoft 97-2003)
            resume_text = extract_text_from_doc_97_2003(doc_file)
        else:
            st.warning(f"Unsupported file type: {doc_file.type}")
            resume_text = ""
    except FileNotFoundError:
        st.warning(f"File not found: {doc_file.name}")
        resume_text = ""

    return resume_text

def extract_text_from_docx(doc_file):
    try:
        doc = Document(doc_file)
        resume_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        st.warning(f"Error extracting text from DOCX file. Error: {e}")
        resume_text = ""

    return resume_text

def extract_text_from_doc_97_2003(doc_file):
    try:
        with open(doc_file.name, 'rb') as file:
            doc = Document(io.BytesIO(file.read()))
            resume_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except FileNotFoundError:
        st.warning(f"File not found: {doc_file.name}")
        resume_text = ""
    except Exception as e:
        st.warning(f"Error extracting text from Word file. Error: {e}")
        resume_text = ""

    return resume_text

def main():
    st.markdown(
        """
        <style>
        .stApp {
            background: linear-gradient(135deg, pink, cyan);
        }
        .classification-section {
            display: inline-block;
            padding: 10px 20px;
            border-radius: 10px;
            background-color: #FFC0CB;
            margin-bottom: 20px;
            text-align: center;
        }
        .classification-section h1 {
            color: black;
            margin: 0;
            padding: 10px 0;
            font-size: 56px;
        }
        </style>
        """,
        unsafe_allow_html=True
    )
    st.markdown(
        "<div class='classification-section'><h1>RESUME CLASSIFICATION</h1></div>",
        unsafe_allow_html=True
    )

    st.markdown("<span style='font-size: 24px; font-weight: bold;'>Upload Your Resume</span>", unsafe_allow_html=True)
    upload_file = st.file_uploader("", type=['docx', 'pdf', 'doc', 'txt'], accept_multiple_files=True)

    if st.button("Process"):
        for doc_file in upload_file:
            if doc_file is not None:
                file_details = {'filename': [doc_file.name],
                                'filetype': doc_file.type.split('.')[-1].upper(),
                                'filesize': str(doc_file.size) + ' KB'}
                file_type = pd.DataFrame(file_details)
                st.write(file_type.set_index('filename'))

                try:
                    is_pdf = doc_file.type == "application/pdf"

                    if is_pdf:
                        with pdfplumber.open(doc_file) as pdf:
                            pages = pdf.pages[0]
                            resume_text = pages.extract_text()
                    else:
                        resume_text = extract_text_from_doc(doc_file)

                    cleaned_text = preprocess(resume_text)

                    predicted_category = loaded_nb_model.predict([cleaned_text])[0]

                    category_names = {0: "PeopleSoft", 1: "React Developer", 2: "SQL Developer", 3: "Workday"}
                    category_name = category_names.get(predicted_category, "Unknown")

                    st.subheader(f"Candidate's Resume matches {category_name} category.")

                except PDFSyntaxError as e:
                    st.warning(f"Unable to process the file {doc_file.name}. Error: {e}")
                except Exception as e:
                    st.warning(f"Error processing document {doc_file.name}. Please ensure it is a valid document. Error: {e}")

if __name__ == '__main__':
    main()
